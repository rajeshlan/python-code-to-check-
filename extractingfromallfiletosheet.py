import os
import re
import openpyxl
import zipfile
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import logging
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for PDF processing

# Configure logging
logging.basicConfig(level=logging.INFO, filename='extraction_log.txt', filemode='w')

# Function to extract details from text
def extract_details(text):
    details = {
        'Name': [],
        'Phone Numbers': [],
        'Contact Numbers': [],
        'Emails': [],
        'Country': [],
        'State': [],
        'City': [],
        'Pin Code': []
    }

    # Compiled regex patterns for efficiency
    name_pattern = re.compile(r'^[A-Za-z\s]+')
    phone_pattern = re.compile(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b')
    contact_pattern = re.compile(r'\+\d+\s\d+[-.\s]?\d+')
    email_pattern = re.compile(r'\S+@\S+')
    # Regex patterns for extracting country, state, city, and pin code
    country_pattern = re.compile(r'Country:\s*([A-Za-z\s]+)')
    state_pattern = re.compile(r'State:\s*([A-Za-z\s]+)')
    city_pattern = re.compile(r'City:\s*([A-Za-z\s]+)')
    pin_code_pattern = re.compile(r'Pin Code:\s*(\d+)')

    # Extract names (assuming the name appears in the first line)
    name_match = name_pattern.search(text)
    if name_match:
        details['Name'].append(name_match.group())

    # Extract phone numbers
    details['Phone Numbers'] = phone_pattern.findall(text)

    # Extract contact numbers with country codes (e.g., +1 1234567890)
    details['Contact Numbers'] = contact_pattern.findall(text)

    # Extract email addresses
    details['Emails'] = email_pattern.findall(text)

    # Extract country, state, city, and pin code
    country_match = country_pattern.search(text)
    if country_match:
        details['Country'].append(country_match.group(1).strip())

    state_match = state_pattern.search(text)
    if state_match:
        details['State'].append(state_match.group(1).strip())

    city_match = city_pattern.search(text)
    if city_match:
        details['City'].append(city_match.group(1).strip())

    pin_code_match = pin_code_pattern.search(text)
    if pin_code_match:
        details['Pin Code'].append(pin_code_match.group(1).strip())

    return details

# Function to extract text from DOCX files
def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    except PackageNotFoundError as e:
        logging.warning(f"Skipping non-Word file: {docx_path}")

    return text

# Function to extract text from images in a DOC file
def extract_text_from_images(doc_path):
    text = ""
    try:
        doc = Document(doc_path)
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                image_path = os.path.join(os.path.dirname(doc_path), doc.part.rels[rel].target_ref)
                image_text = extract_text_from_image(image_path)
                text += image_text + '\n'
    except Exception as e:
        logging.warning(f"Error extracting text from images: {str(e)}")
    return text

# Function to extract text from an image using pytesseract
def extract_text_from_image(image_path):
    text = ""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
    except Exception as e:
        logging.warning(f"Error extracting text from image: {str(e)}")
    return text

# Function to extract text from PDF files
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
    except Exception as e:
        logging.warning(f"Error extracting text from PDF: {str(e)}")
    return text

# Function to process files in a folder
def process_files(input_folder, output_file):
    # Create a new Excel workbook and add a worksheet
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    # Set the headers for the columns
    headers = ['File Name', 'Name', 'Phone Numbers', 'Contact Numbers', 'Emails', 'Country', 'State', 'City', 'Pin Code']
    worksheet.append(headers)

    for filename in os.listdir(input_folder):
        if filename.endswith(('.txt', '.docx', '.doc', '.pdf')):
            file_path = os.path.join(input_folder, filename)

            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif filename.endswith('.docx') or filename.endswith('.doc'):
                try:
                    # Attempt to extract text from DOCX files
                    text = extract_text_from_docx(file_path)
                    # Extract text from images in DOC file
                    image_text = extract_text_from_images(file_path)
                    text += image_text
                except (KeyError, zipfile.BadZipFile) as e:
                    logging.error(f"Error processing file: {file_path}")
                    continue  # Skip this file and proceed to the next one
                except Exception as e:
                    logging.warning(f"Skipping non-Word file: {file_path}")
                    continue
            elif filename.endswith('.pdf'):
                try:
                    # Attempt to extract text from PDF files
                    text = extract_text_from_pdf(file_path)
                except Exception as e:
                    logging.warning(f"Error processing PDF file: {file_path}")
                    continue

            details = extract_details(text)

            # Create a list of values to be added to the worksheet
            values = [filename]
            for key in headers[1:]:
                values.append(', '.join(details.get(key, [''])))  # Use a default empty string if key not found

            # Add the values to the worksheet
            worksheet.append(values)

    # Save the workbook to a file
    workbook.save(output_file)

    logging.info(f"Extraction completed. Results saved to '{output_file}'")

if __name__ == "__main__":
    # Input directory containing text, DOCX files, images in DOC files, and PDF files
    #input_folder = r'C:\Users\rajes\OneDrive\Desktop\CV SEPT112023-20230913T044719Z-001\CV SEPT112023'
    input_folder = r'C:\Users\rajes\OneDrive\Desktop\CV (1)'
    output_file = 'extracted_details_all_details.xlsx'

    process_files(input_folder, output_file)
